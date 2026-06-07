import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_global_underfill_summary():
    doc = Document()
    
    # Font Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

    # --- Title Page ---
    title = doc.add_heading('Underfill (底填膠) 技術全球通論與產業標準白皮書', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Global Technology Overview, Industry Guidelines & Reliability Standards').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n')

    # --- 1. 技術基礎與物理原理 (Fundamentals) ---
    doc.add_heading('1. 技術基礎與物理原理', level=1)
    doc.add_paragraph('底填技術的核心在於解決半導體封裝中「熱膨脹係數 (CTE) 失配」的問題。')
    doc.add_paragraph('主要物理目標：', style='Normal').bold = True
    bullets = [
        '應力緩衝：建立矽晶圓 (CTE ~2.6 ppm) 與有機基板 (CTE ~17 ppm) 之間的剛性連接，減少錫球焊點的剪切應力。',
        '毛細作用 (Capillary Action)：利用液體在窄縫中的表面張力自動填充，受 Washburn 方程支配。',
        '環境保護：封裝焊點以防止濕氣、化學污染及離子遷移。'
    ]
    for b in bullets: doc.add_paragraph(b, style='List Bullet')

    # --- 2. 產業主流技術分類 (Comparison Table) ---
    doc.add_heading('2. 產業主流技術分類', level=1)
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ['技術類別', '工藝時機', '核心優點', '典型應用']
    for i, h in enumerate(headers): table.rows[0].cells[i].text = h
    
    data = [
        ('毛細型 (CUF)', '回流焊後', '高可靠度、低 CTE、工藝成熟', '手機、數據中心、車用'),
        ('非流動型 (NFU)', '回流焊前', '產能高、省略點膠固化步驟', '低 I/O 消費電子'),
        ('壓塑型 (MUF)', '模封過程中', '極高產能、多晶片同時封裝', '存儲晶片、大型 BGA'),
        ('晶圓級 (WLUF)', '晶圓切割前', '適合極微細間隙與高密度組裝', '先進封裝 (2.5D/3D)')
    ]
    for i, row in enumerate(data):
        for j, val in enumerate(row): table.rows[i+1].cells[j].text = val

    # --- 3. 材料科學關鍵指標 (Material Science) ---
    doc.add_heading('3. 材料科學關鍵指標', level=1)
    doc.add_paragraph('選擇底填膠時，必須平衡流動性與機械性能。')
    
    mat_table = doc.add_table(rows=6, cols=2)
    mat_table.style = 'Table Grid'
    m_specs = [
        ('CTE (熱膨脹係數)', '必須盡可能降低以匹配矽膠 (Silica Fillers 比例通常 > 60%)。'),
        ('Tg (玻璃轉化溫度)', '決定材料穩定工作的最高度。Tg 以上 CTE 會急劇上升 (CTE 2)。'),
        ('Modulus (彈性模量)', '決定材料支撐焊點防止熱變形的剛性。'),
        ('Viscosity (粘度)', '決定填充速度，對於 <50um 的間隙，需極低粘度。'),
        ('Adhesion (附著力)', '與阻焊漆 (Solder Mask) 及晶片鈍化層的化學鍵結能力。')
    ]
    for i, (k, v) in enumerate(m_specs):
        mat_table.rows[i].cells[0].text = k
        mat_table.rows[i].cells[1].text = v

    # --- 4. 全球產業標準 (Industry Guidelines) ---
    doc.add_heading('4. 全球產業標準與規範', level=1)
    ipc_bullets = [
        'IPC J-STD-030：底填材料選擇與工藝應用的權威指導書。',
        'IPC-A-610：針對終端組件目視檢查的全球標準。',
        'JEDEC Standard：針對溫度循環 (TCT) 與落下測試 (Drop Test) 的規範。'
    ]
    for b in ipc_bullets: doc.add_paragraph(b, style='List Bullet')

    # --- 5. 故障分析與可靠性 (Failure Analysis) ---
    doc.add_heading('5. 常見失效模式與可靠性分析', level=1)
    fail_table = doc.add_table(rows=6, cols=3)
    fail_table.style = 'Table Grid'
    f_h = ['失效現象', '物理原因', '後續風險']
    for i, h in enumerate(f_h): fail_table.rows[0].cells[i].text = h
    
    f_data = [
        ('分層 (Delamination)', '界面應力大於附著力', '水汽入侵、錫鬚生長、焊點斷裂'),
        ('內部空洞 (Voids)', '氣體受困或有機物出氣', '應力集中、熱疲勞加速'),
        ('爆米花效應', '受潮膠材在迴焊時水汽急劇膨脹', '封裝開裂、元件失效'),
        ('錫球擠出 (Extrusion)', '分層空隙導致錫球受熱溢流', '導電短路、燒毀'),
        ('膠材龜裂 (Cracking)', '應力分布不均或 Tg 過低', '可靠度降低')
    ]
    for i, row in enumerate(f_data):
        for j, val in enumerate(row): fail_table.rows[i+1].cells[j].text = val

    # --- 6. 進階檢測技術 (Advanced Inspection) ---
    doc.add_heading('6. 進階檢測工法與技術', level=1)
    doc.add_paragraph('由於底填膠被遮蓋，傳統目檢失效，需依賴：')
    inspect = [
        'CSAM (超音波掃描)：偵測界面分層與微小空洞的最靈敏手段。',
        '3D X-Ray：觀察焊點形狀與底填膠大規模填充異常。',
        'Cross-Section (切片)：破壞性檢測，用於精確觀察填充高度與界面微觀結構。'
    ]
    for b in inspect: doc.add_paragraph(b, style='List Bullet')

    # --- 7. 未來趨勢 (Future Trends) ---
    doc.add_heading('7. 先進封裝趨勢', level=1)
    doc.add_paragraph('隨著 AI 與 HPC (高效能運算) 發展，底填技術正面臨：')
    trends = [
        '高導熱底填 (Thermal Underfill)：具備散熱路徑功能的材料。',
        '低介電常數 (Low-k) 兼容：保護脆弱的 Low-k 結構不被剝離。',
        '2.5D/3D 堆疊填充：極小隙縫 (Gap < 10um) 的真空輔助填充技術。'
    ]
    for b in trends: doc.add_paragraph(b, style='List Bullet')

    save_path = 'Underfill_Global_Tech_Summary_Whitepaper.docx'
    doc.save(save_path)
    print(f"Global Tech Summary saved to: {os.path.abspath(save_path)}")

if __name__ == '__main__':
    create_global_underfill_summary()

