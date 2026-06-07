import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_enhanced_wi():
    doc = Document()
    
    # Font Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

    # --- Title Section ---
    title = doc.add_heading('Underfill 管理與作業標準書 (Work Instruction)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table_info = doc.add_table(rows=4, cols=4)
    table_info.style = 'Table Grid'
    info = [
        ['文件編號', 'WI-SMT-UF3808-PRO', '文件版本', 'V2.0'],
        ['膠材型號', 'Loctite ECCOBOND UF 3808', '存儲溫度', '-15°C to -40°C'],
        ['制訂部門', 'SMT 工程部', '生效日期', '2026-04-09'],
        ['機密等級', '內部公開', '審核', 'Approved']
    ]
    for i, row_data in enumerate(info):
        for j, val in enumerate(row_data):
            table_info.rows[i].cells[j].text = val

    doc.add_paragraph('\n')

    # --- 1. 目的 ---
    doc.add_heading('1. 目的', level=1)
    doc.add_paragraph('本文件旨在詳細規範 Loctite ECCOBOND UF 3808 在自動化與半自動化點膠製程中的技術參數。透過標準化 Ball Pitch、Ball Size 及 Die Size 的對應關係，確保焊點在極端溫差與機械應力下的穩定性，並符合 IPC 最高等級之電子組件可靠度標準。')

    # --- 2. 膠材原廠規範與深度特性 (Henkel Loctite ECCOBOND UF 3808) ---
    doc.add_heading('2. 膠材原廠規範與物理特性', level=1)
    doc.add_paragraph('Loctite UF 3808 是一款專為「高生產率」與「可維修性」設計的單組份底填膠。')
    
    table_specs = doc.add_table(rows=12, cols=2)
    table_specs.style = 'Table Grid'
    specs = [
        ('化學基礎 (Chemical Base)', '環氧樹脂 (One-component Epoxy)'),
        ('顏色 (Color)', '黑色 (Black)'),
        ('粘度 (Viscosity @ 25°C)', '348 mPa·s (cP) @ 1,000 s⁻¹ - 具備極佳浸潤性'),
        ('儲存週期 (Shelf Life)', '12 個月 (@ -40°C 至 -15°C)'),
        ('室溫工作壽命 (Pot Life)', '24 小時 (@ 25°C)'),
        ('玻璃轉化溫度 (Tg)', '113°C (by TMA)'),
        ('熱膨脹係數 (CTE) < Tg', '55 ppm/°C'),
        ('熱膨脹係數 (CTE) > Tg', '171 ppm/°C'),
        ('儲能模量 (Storage Modulus)', '2.8 GPa (@ 25°C)'),
        ('硬度 (Hardness, Shore D)', '88'),
        ('固化推薦 (Recommend Cure)', '130°C / 8 mins 或 150°C / 5 mins'),
        ('化學抗性 (Chemical Resistance)', '優異的抗溶劑性與低吸濕率，可通過鉛封與鹽霧測試')
    ]
    for i, (k, v) in enumerate(specs):
        table_specs.rows[i].cells[0].text = k
        table_specs.rows[i].cells[1].text = v

    # --- 3. 應用元件類別矩陣對應表 (尺寸與間距詳述) ---
    doc.add_heading('3. 應用元件類別與物理尺寸矩陣表', level=1)
    doc.add_paragraph('本表定義了膠材流動性與元件幾何尺寸的兼容性基準。')
    
    table_matrix = doc.add_table(rows=6, cols=5)
    table_matrix.style = 'Table Grid'
    headers = ['元件類別', 'Die Size (mm)', 'Ball Pitch (mm)', 'Ball Size (mm)', '適用評估']
    for i, h in enumerate(headers):
        table_matrix.rows[0].cells[i].text = h
    
    matrix_data = [
        ('Standard BGA', '15x15 至 45x45', '0.50 至 1.27', '0.30 至 0.60', '建議使用 (增強衝擊力)'),
        ('Fine Pitch BGA', '5x5 至 15x15', '0.35 至 0.40', '0.20 至 0.25', '建議使用 (防止熱裂紋)'),
        ('WLCSP / CSP', '1.5x1.5 至 8x8', '0.30 至 0.50', '0.15 至 0.25', '必須使用 (高風險件)'),
        ('Flip Chip', '2x2 至 12x12', '0.10 至 0.25', '0.05 至 0.15', '核心應用 (底部隙 30-80um)'),
        ('QFN / LGA', '< 10x10', '0.40 至 0.65', 'N/A (Pad)', '條件適用 (需確認 Standoff)')
    ]
    for i, row in enumerate(matrix_data):
        for j, val in enumerate(row):
            table_matrix.rows[i+1].cells[j].text = val

    # --- 4. 品質檢驗標準與 IPC 規範細則 ---
    doc.add_heading('4. 品質檢驗標準與 IPC 規範細則', level=1)
    
    doc.add_heading('4.1 IPC 規範具體條文參考', level=2)
    ipc_table = doc.add_table(rows=5, cols=2)
    ipc_table.style = 'Table Grid'
    ipc_content = [
        ('IPC-J-STD-030 Section 4.2', '定義膠材流動速度 (Flow Rate) 測試：必須確保在特定溫度下填滿元件底部的時間符合工藝預期。'),
        ('IPC-A-610 Section 8.3.1', '定義底部填充應覆蓋元件周長的 100%，且溢膠邊寬 (Fillet width) 不得小於元件邊緣。'),
        ('IPC-7095 Section 7.5.3', '空洞判定標準：總空洞面積不可超過焊點接觸面積的 25%。單個大空洞不得跨越相鄰兩個焊點。'),
        ('IPC-7351', '焊盤設計規範：元件周圍必須保留至少 1.5mm 的點膠避讓區，防止污染鄰近元件或測試點。')
    ]
    for i, (k, v) in enumerate(ipc_content):
        ipc_table.rows[i].cells[0].text = k
        ipc_table.rows[i].cells[1].text = v

    doc.add_heading('4.2 具體允收/拒收標準 (Criteria)', level=2)
    criteria = [
        '1. 包封高度 (Fillet Height): 理想值為元件高度的 50%-75%。拒收：若包封高度超過元件頂部 (Top Surface)，視為污染。',
        '2. 溢流區域 (Bleed-out): 膠材邊緣擴散不得超過元件周邊 1.0mm 以上。拒收：若膠材侵入相鄰元件 Pad 或測試針點。',
        '3. 表面缺陷: 允收：Fillet 表面平滑且無肉眼可見的大型氣泡。拒收：Fillet 存在裂縫或嚴重斷開 (Void at edge)。',
        '4. 內部空洞 (Voids by X-Ray): 所有錫球下方的空洞總合須 < 總體積的 25%。拒收：關鍵信號 (Signal Pins) 下方有貫穿性空洞。',
        '5. 元件傾斜: 允收：元件水平傾斜度 < 5 度。拒收：膠材不均導致元件單側抬升。'
    ]
    for c in criteria:
        doc.add_paragraph(c, style='List Bullet')

    # --- 5. 作業注意事項與 Rework ---
    doc.add_heading('5. 作業注意事項與返修規範', level=1)
    p = doc.add_paragraph()
    p.add_run('UF 3808 具備良好的可修復性 (Reworkable)。').bold = True
    p.add_run('\n返修條件：')
    p.add_run('\n- 設備：熱風 BGA 返修台')
    p.add_run('\n- 溫度：組件需加熱至 230°C - 245°C (Reflow Peak temp)')
    p.add_run('\n- 拆卸：使用專用不銹鋼刮刀在軟化時緩慢移除殘膠，嚴禁暴力撬起。')
    p.add_run('\n- 清理：使用異丙基酒精 (IPA) 或原廠推薦清洗劑進行殘存環氧樹脂清理。')

    save_path = 'Underfill_WI_V2_Henkel_Standard.docx'
    doc.save(save_path)
    print(f'Enhanced Document saved to: {os.path.abspath(save_path)}')

if __name__ == '__main__':
    create_enhanced_wi()
