import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_rich_wi():
    doc = Document()
    
    # Paths for images (Update these to the actual absolute paths from the previous turn)
    IMG_PATTERNS = r'C:\Users\user\.gemini\antigravity\brain\db344bb8-a2bb-486c-81a4-f44245a98563\underfill_dispensing_patterns_1775741805036.png'
    IMG_FILLET = r'C:\Users\user\.gemini\antigravity\brain\db344bb8-a2bb-486c-81a4-f44245a98563\underfill_fillet_height_1775741826978.png'
    IMG_VOIDS = r'C:\Users\user\.gemini\antigravity\brain\db344bb8-a2bb-486c-81a4-f44245a98563\underfill_voiding_xray_1775741846746.png'

    # Global Font Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

    # --- Title Section ---
    title = doc.add_heading('Underfill 全工藝作業標準說明書 (圖文強化版)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')

    # --- 1. 目的與適用範圍 ---
    doc.add_heading('1. 目的與適用範圍', level=1)
    doc.add_paragraph('本文件定義 Loctite UF 3808 之封裝底填工藝標準，旨在透過視覺化圖示與嚴格參數，確保產線人員能正確執行點膠與品質判定。')

    # --- 2. 膠材技術參數 ---
    doc.add_heading('2. 膠材技術參數 (Loctite UF 3808)', level=1)
    table_specs = doc.add_table(rows=6, cols=2)
    table_specs.style = 'Table Grid'
    specs = [
        ('粘度 (Viscosity)', '348 cP @ 25°C - 高流動性'),
        ('固化條件', '130°C / 8 mins 或 150°C / 5 mins'),
        ('Tg / CTE', '113°C / 55-171 ppm'),
        ('儲存週期', '12 個月 (@ -40°C 至 -15°C)'),
        ('工作壽命', '24 小時 (@ 25°C)'),
        ('優點', '高附著力、優異抗震性、可返修')
    ]
    for i, (k, v) in enumerate(specs):
        table_specs.rows[i].cells[0].text = k
        table_specs.rows[i].cells[1].text = v

    # --- 3. 點膠路徑設計 ---
    doc.add_heading('3. 點膠路徑與模式示意圖', level=1)
    doc.add_paragraph('正確的路徑是確保無空洞填滿的核心。推薦優先使用 L 型點膠。')
    
    if os.path.exists(IMG_PATTERNS):
        doc.add_picture(IMG_PATTERNS, width=Inches(4.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('圖 1: 點膠模式示意圖 (推薦 L 型路徑)', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph('[圖片預留區: 點膠路徑示意圖]', style='Normal')

    # --- 4. 應用元件類別矩陣 ---
    doc.add_heading('4. 應用元件矩陣表 (尺寸/間距)', level=1)
    table_matrix = doc.add_table(rows=5, cols=5)
    table_matrix.style = 'Table Grid'
    h = ['元件', 'Die Size', 'Pitch', 'Ball Size', '判定']
    for i, v in enumerate(h): table_matrix.rows[0].cells[i].text = v
    m_data = [
        ('BGA', '> 25mm', '0.8~1.27', '0.45~0.60', '推薦'),
        ('WLCSP', '< 10mm', '0.3~0.5', '0.15~0.25', '強制'),
        ('Flip Chip', '2~15mm', '0.1~0.25', '0.05~0.15', '核心'),
        ('QFN', '< 12mm', '0.4~0.65', 'N/A', '評估')
    ]
    for i, row in enumerate(m_data):
        for j, val in enumerate(row): table_matrix.rows[i+1].cells[j].text = val

    # --- 5. 品質檢驗標準 (圖文對照) ---
    doc.add_heading('5. 品質檢驗與 IPC 允收標準', level=1)
    
    doc.add_heading('5.1 Fillet (膠側包封) 高度要求', level=2)
    doc.add_paragraph('依據 IPC-A-610，膠側包封高度需介於元件厚度的 50% 至 75% 之間。')
    
    if os.path.exists(IMG_FILLET):
        doc.add_picture(IMG_FILLET, width=Inches(4.0))
        doc.add_paragraph('圖 2: Fillet 高度合格標準示意圖 (50%-75% Height)', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('5.2 內部空洞 (Voiding) 判定', level=2)
    doc.add_paragraph('利用 X-Ray 進行非破壞性檢驗。總空洞面積不可超過填膠區域的 25%。')
    
    if os.path.exists(IMG_VOIDS):
        doc.add_picture(IMG_VOIDS, width=Inches(4.0))
        doc.add_paragraph('圖 3: X-Ray 下空洞判定 (Voids < 25%)', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 6. 異常對策表 ---
    doc.add_heading('6. 常見異常原因與對策', level=1)
    table_ts = doc.add_table(rows=4, cols=3)
    table_ts.style = 'Table Grid'
    ts_data = [
        ('異常現象', '主因', '對策'),
        ('滲透不全', '基板預熱不足', '提升預熱至 85°C'),
        ('膠材分層', 'PCB 髒污', '強化回流後清洗'),
        ('表面氣泡', '回溫不足', '落實 2H 回溫')
    ]
    for i, row in enumerate(ts_data):
        for j, val in enumerate(row): table_ts.rows[i].cells[j].text = val

    # --- 7. 工廠現場照片預留 ---
    doc.add_heading('7. 工廠現場作業參考照片 (預留)', level=1)
    doc.add_paragraph('[請工程人員在此插入實機操作照片: 針頭設定、泵壓控管、烘箱 Profile 控制圖表]', style='Normal')

    save_path = 'Underfill_WI_Graphic_Rich.docx'
    doc.save(save_path)
    print(f'Rich Graphic Document saved to: {os.path.abspath(save_path)}')

if __name__ == '__main__':
    create_rich_wi()
