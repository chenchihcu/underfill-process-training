from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def create_master_handbook(out_path, media_dir):
    # Standard document setup
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

    # Image Paths
    IMG = {
        'PATTERNS': r'C:\Users\user\.gemini\antigravity\brain\db344bb8-a2bb-486c-81a4-f44245a98563\underfill_dispensing_patterns_1775741805036.png',
        'FILLET': r'C:\Users\user\.gemini\antigravity\brain\db344bb8-a2bb-486c-81a4-f44245a98563\underfill_fillet_height_1775741826978.png',
        'XRAY': r'C:\Users\user\.gemini\antigravity\brain\db344bb8-a2bb-486c-81a4-f44245a98563\underfill_voiding_xray_1775741846746.png',
        'CSAM': r'C:\Users\user\.gemini\antigravity\brain\db344bb8-a2bb-486c-81a4-f44245a98563\csam_underfill_inspection_1775742090770.png',
        'NEEDLE': r'C:\Users\user\.gemini\antigravity\brain\db344bb8-a2bb-486c-81a4-f44245a98563\underfill_needle_setup_1775742113973.png'
    }

    # --- COVER PAGE ---
    doc.add_heading('Underfill Technology Master Handbook', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('【究極整合版：工藝理論、材料規格與作業指導】').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Industrial White Paper + Material TDS + Standard Operating Procedure').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n' * 2)

    # --- PART 1: 全球底填技術白皮書 (GLOBAL WHITE PAPER) ---
    doc.add_heading('PART 1: 全球底填技術概論與物理原理', level=1)
    doc.add_heading('1.1 基本物理目標', level=2)
    bullets = [
        'CTE Mismatch 應力管理：緩解矽晶圓(2.6ppm)與基板(17ppm)間的熱脹冷縮失配。',
        '毛細流動物理 (Capillary Science)：受 Washburn 方程支配的滲透行為。',
        '結構強化：提升焊點疲勞壽命與機械抗衝擊能力 (Drop Test)。'
    ]
    for b in bullets: doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('1.2 產業主流技術分類對照', level=2)
    table_tech = doc.add_table(rows=5, cols=4); table_tech.style = 'Table Grid'
    h_tech = ['類別', '時機', '優點', '限制']
    for i,h in enumerate(h_tech): table_tech.rows[0].cells[i].text = h
    tech_data = [
        ('CUF (毛細型)', '回流焊後', '高可靠度、低 CTE、彈性高', '步驟較慢'),
        ('NFU (非流動)', '回流焊前', '高產能、設備成本低', '有限填充比例'),
        ('MUF (壓塑型)', '模封過程', '極高產能、大規模生產', '模具成本高'),
        ('WLUF (晶圓級)', '切割前', '微小間隙與先進封裝首選', '工藝複雜度高')
    ]
    for i,row in enumerate(tech_data):
        for j,v in enumerate(row): table_tech.rows[i+1].cells[j].text = v

    # --- PART 2: 材料規格與性質 (MATERIAL SPECS - LOCTITE UF 3808) ---
    doc.add_heading('PART 2: Loctite ECCOBOND UF 3808 深度材料規格', level=1)
    doc.add_paragraph('本章節整合 Henkel 原廠最新物理與化學技術指標：')
    table_m = doc.add_table(rows=12, cols=2); table_m.style = 'Table Grid'
    specs = [
        ('化學類型', '單組份環氧樹脂熱固化'),
        ('粘度 (Viscosity)', '348 mPa·s (cP) @ 25°C, 1,000 s⁻¹'),
        ('固化推薦', '130°C / 8 mins 或 150°C / 5 mins'),
        ('Tg (玻璃轉化點)', '113°C (by TMA)'),
        ('CTE 1 (<Tg)', '55 ppm/°C'),
        ('CTE 2 (>Tg)', '171 ppm/°C'),
        ('儲能模量 (Modulus)', '2.8 GPa (@ 25°C)'),
        ('硬度 (Shore D)', '88'),
        ('存儲壽命 (Shelf)', '12 個月 (@ -15°C 至 -40°C)'),
        ('工作時間 (Pot Life)', '24 小時 (@ 25°C)'),
        ('吸濕率', '< 0.5%'),
        ('化學抵抗', '具備優異的抗溶劑、抗助焊劑殘留性能')
    ]
    for i, (k, v) in enumerate(specs):
        table_m.rows[i].cells[0].text = k
        table_m.rows[i].cells[1].text = v

    # --- PART 3: 作業標準指導 (STANDARD OPERATING PROCEDURE) ---
    doc.add_heading('PART 3: 底填膠標準作業流程與參數指導', level=1)
    doc.add_heading('3.1 生產前置準備 (Preparation)', level=2)
    prep = [
        '進料檢驗 (IQC)：查核 CoA 報告、粘度與批次日期。',
        '冷凍存儲：-15C 以下。取出後針頭朝下垂直靜置 2 小時進行回溫。',
        'PCB 預烤：100°C / 1.5H。確保無水汽殘留在過孔或基板層中。',
        '基板預熱：點膠平台須預熱至 75°C - 90°C，以利膠材迅速鋪展。'
    ]
    for p in prep: doc.add_paragraph(p, style='List Bullet')

    doc.add_heading('3.2 點膠參數設計與路徑 (Engineering Setup)', level=2)
    if os.path.exists(IMG['NEEDLE']):
        doc.add_picture(IMG['NEEDLE'], width=Inches(4.5))
        doc.add_paragraph('圖 1: 針頭設定指標 (23G-27G, 高度 0.3mm)', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(IMG['PATTERNS']):
        doc.add_picture(IMG['PATTERNS'], width=Inches(4.5))
        doc.add_paragraph('圖 2: L型推薦路徑圖解 (引導氣體排出最佳效果)', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_params = [
        '點膠壓力：0.15 - 0.35 MPa。',
        '點膠延遲 (Delay)：150 - 500 ms (防止拉絲污染元件表面)。',
        '點膠高度：距離元件邊緣 0.5 - 0.8 mm。'
    ]
    for pp in p_params: doc.add_paragraph(pp, style='List Bullet')

    doc.add_heading('3.3 應用元件尺寸矩陣 (Applicability Matrix)', level=2)
    table_mx = doc.add_table(rows=6, cols=5); table_mx.style = 'Table Grid'
    mx_h = ['元件', 'Die Size', 'Pitch', 'Ball Size', '判定']
    for i,h in enumerate(mx_h): table_mx.rows[0].cells[i].text = h
    mx_data = [
        ('Flip Chip', '2~15mm', '0.1~0.25mm', '0.05~0.15mm', '核心對象'),
        ('CSP/WLCSP', '1.5~10mm', '0.3~0.5mm', '0.15~0.25mm', '強制使用'),
        ('Small BGA', '5~15mm', '0.35~0.5mm', '0.2~0.3mm', '高度建議'),
        ('Large BGA', '15~45mm', '0.8~1.27mm', '0.35~0.6mm', '建議使用'),
        ('QFN/LGA', '< 12mm', '0.4~0.65mm', 'Pad Only', '評估 Standoff')
    ]
    for i,row in enumerate(mx_data):
        for j,v in enumerate(row): table_mx.rows[i+1].cells[j].text = v

    # --- PART 4: 品質檢驗標準 (QUALITY & IPC) ---
    doc.add_heading('PART 4: 品質檢驗允收標準 (IPC Standard)', level=1)
    doc.add_heading('4.1 IPC 規範具體對照 (Clause Mapping)', level=2)
    table_ipc = doc.add_table(rows=4, cols=2); table_ipc.style = 'Table Grid'
    ipc_d = [
        ('IPC-J-STD-030 4.2', '定義填充流速、浸潤能力與固化時間標準。'),
        ('IPC-A-610 8.3.1', '定義 Fillet 週長覆蓋須 100%，高度須覆蓋側壁 50%-75%。'),
        ('IPC-7095 7.5.3', '定義 X-Ray 空洞判定：單一焊點或總面積空洞均須 < 25%。')
    ]
    for i,(k,v) in enumerate(ipc_d):
        table_ipc.rows[i].cells[0].text = k
        table_ipc.rows[i].cells[1].text = v

    doc.add_heading('4.2 視覺化判定準則與檢測', level=2)
    if os.path.exists(IMG['FILLET']):
        doc.add_picture(IMG['FILLET'], width=Inches(4.0))
        doc.add_paragraph('圖 3: Fillet 高度合格對比圖', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(IMG['XRAY']):
        doc.add_picture(IMG['XRAY'], width=Inches(3.5))
        doc.add_paragraph('圖 4: X-Ray 內部空洞判定', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(IMG['CSAM']):
        doc.add_picture(IMG['CSAM'], width=Inches(4.0))
        doc.add_paragraph('圖 5: CSAM 分層與剝離分析', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- PART 5: 異常分析與故障排除 (TROUBLESHOOTING) ---
    doc.add_heading('PART 5: 異常分析與故障排除對策矩陣', level=1)
    table_ts = doc.add_table(rows=6, cols=3); table_ts.style = 'Table Grid'
    ts_h = ['失效現象', '根本原因 (Root Cause)', '改善對策 (Action)']
    for i,h in enumerate(ts_h): table_ts.rows[0].cells[i].text = h
    ts_data = [
        ('滲透不全', '基板溫度低、助焊劑殘留、間隙過窄', '校準基板預熱 85C、強化洗淨工程。'),
        ('微氣泡/空洞', 'PCB 受潮、膠材回溫不足、點膠路徑錯誤', 'PCB 100C 預烤、嚴格執行 2H 垂直回溫。'),
        ('溢流污染', '點膠壓力大、元件 Standoff 太低', '調降點膠出膠量、修正點膠高度。'),
        ('爆米花效應', '濕氣急劇膨脹', '控管 MSD 等級與烘烤流程。'),
        ('界面分層', '表面能過低或清潔不徹底', '調整 Plasma 處理參數或優化洗劑。')
    ]
    for i,row in enumerate(ts_data):
        for j,v in enumerate(row): table_ts.rows[i+1].cells[j].text = v

    # --- PART 6: 返修程序 (REWORK) ---
    doc.add_heading('PART 6: 受控返修標準程序 (Rework Procedure)', level=1)
    rework = [
        '步驟 1 (Pre-heat)：全板預熱至 150°C 穩定組件熱應力。',
        '步驟 2 (Localized)：噴嘴加熱至 240-250°C 弱化環氧樹脂結構。',
        '步驟 3 (Removal)：使用真空吸嘴垂直移除元件。',
        '步驟 4 (Clean)：在餘熱下使用不銹鋼或聚四氟乙烯刮刀清理焊盤。',
        '步驟 5 (Final)：使用 IPA 酒精進行最終界面清洗。'
    ]
    for step in rework: doc.add_paragraph(step, style='List Bullet')

    # --- PART 7: 操作員檢核表 (CHECKLIST) ---
    doc.add_heading('PART 7: 操作員每日班次檢核表', level=1)
    chk = [
        '[ ] 批次有效期確認 (Expiry Date) ',
        '[ ] 冷凍櫃溫度查核 (-15°C ~ -40°C)',
        '[ ] 垂直回溫紀錄 (2.0 Hours min)',
        '[ ] PCB 預熱溫度實測 (75°C - 90°C)',
        '[ ] 針頭清潔度與斷點觀察',
        '[ ] 首件 (First Article) Fillet 與 X-ray 抽測'
    ]
    for item in chk: doc.add_paragraph(item, style='List Bullet')

    # --- SAVE ---
    save_path = 'Underfill_TRUE_ULTIMATE_MASTER_HANDBOOK.docx'
    doc.save(save_path)
    print(f"Master Handbook saved to: {os.path.abspath(save_path)}")

if __name__ == '__main__':
    create_master_handbook(
        r'C:\Users\user\Desktop\Underfill\Underfill_TRUE_ULTIMATE_MASTER_HANDBOOK.docx',
        r'C:\Users\user\Desktop\Underfill\extracted_images'
    )
