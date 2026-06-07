import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def build_true_ultimate_wi(out_path, media_dir):
    doc = Document()
    
    # Font Setup for Professional Look
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

    # Image Mapping
    IMG = {
        'PATTERNS': os.path.join(media_dir, 'word/media/image1.jpeg'),
        'FILLET': os.path.join(media_dir, 'word/media/image2.jpeg'),
        'XRAY': os.path.join(media_dir, 'word/media/image3.jpeg'),
        'CSAM': os.path.join(media_dir, 'word/media/image4.jpeg'),
        'NEEDLE': os.path.join(media_dir, 'word/media/image5.jpeg')
    }

    # --- Header Control Table ---
    title = doc.add_heading('Underfill (底填膠) 全工藝標準作業指導書 (Ultimate SOP)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    hdr_table = doc.add_table(rows=5, cols=4)
    hdr_table.style = 'Table Grid'
    hdr_rows = [
        ['文件編號', 'WI-SMT-UF3808-V6', '文件版本', 'V6.0 (Ultimate)'],
        ['膠材型號', 'Loctite ECCOBOND UF 3808', '存儲規範', '-15C to -40C'],
        ['制訂部門', 'SMT 工程部 / QA 品質部', '生效日期', '2026-04-10'],
        ['應用類別', 'BGA / CSP / Flip Chip', '機密等級', '內部限制'],
        ['核准主管', 'Engineering Director', '頁碼', '1 / 1']
    ]
    for i, r in enumerate(hdr_rows):
        for j, v in enumerate(r): hdr_table.rows[i].cells[j].text = v

    # --- 1. 目的與範圍 (Purpose & Scope) ---
    doc.add_heading('1. 目的與適用範圍', level=1)
    doc.add_paragraph('本文件整合 Henkel 原廠 TDS 規格、IPC 國際通用組裝標準與工業界最佳實踐。旨在提供 SMT 操作員與工程師關於底填膠從進料檢驗(IQC)、存儲、回溫、點膠參數設計、固化後品質檢驗至返修(Rework)的全生命週期標準化指導，確保電子組件具備極高的可靠度與抗應力能力。')

    # --- 2. 膠材規格與物理性質 (Detailed Specs) ---
    doc.add_heading('2. 膠材規格與物理性質 (Loctite UF 3808)', level=1)
    doc.add_paragraph('UF 3808 是業界領先的單組份環氧樹脂底填膠，專為移動設備與高密度電路設計。')
    
    spec_table = doc.add_table(rows=12, cols=2)
    spec_table.style = 'Table Grid'
    spec_data = [
        ('化學基礎 (Chemical Base)', '環氧樹脂 (One-component Epoxy)'),
        ('外觀 / 顏色', '黑色液體 (Black Liquid)'),
        ('粘度 (Viscosity @ 25°C)', '348 mPa·s (cP) @ 1,000 s⁻¹ (極佳的快速滲透能力)'),
        ('玻璃轉化溫度 (Tg)', '113°C (by TMA) - 提供優異的熱循環彈性'),
        ('熱膨脹係數 (CTE) < Tg', '55 ppm/°C'),
        ('熱膨脹係數 (CTE) > Tg', '171 ppm/°C'),
        ('儲能模量 (Storage Modulus)', '2.8 GPa (@ 25°C) - 作為有效的應力緩衝墊'),
        ('硬度 (Shore D Hardness)', '88'),
        ('儲存週期 (Shelf Life)', '12 個月 (@ -15°C 至 -40°C)'),
        ('工作壽命 (Pot Life)', '24 小時 (@ 25°C) - 嚴禁使用超過此時限的文件'),
        ('固化推薦條件', '130°C / 8 mins 或 150°C / 5 mins (板面實測溫度)'),
        ('吸濕率 (Moisture Abs)', '< 0.5% (防止極端環境下的爆板風險)')
    ]
    for i, (k, v) in enumerate(spec_data):
        spec_table.rows[i].cells[0].text = k
        spec_table.rows[i].cells[1].text = v

    # --- 3. 國際規範與品質映射 (IPC Standards) ---
    doc.add_heading('3. 國際規範與品質檢驗標準 (IPC Mapping)', level=1)
    ipc_table = doc.add_table(rows=5, cols=2)
    ipc_table.style = 'Table Grid'
    ipc_mapping = [
        ('IPC-J-STD-030 Section 4.2', '定義底填膠的流動速率測試方法與膠材選擇標準。'),
        ('IPC-A-610 Section 8.3.1', '定義 Fillet (包封) 包覆率須達 100% 周長且覆蓋元件側壁 50%-75%。'),
        ('IPC-7095 Section 7.5.3', '規範 X-Ray 檢測標準：單一焊點下空洞面積不得超過 25%；總空洞率 < 25%。'),
        ('IPC-7711/7721', '規範底填膠元件的熱拆卸流程，加熱設備溫度通常需達 240°C 以上。'),
        ('IPC-SM-785', '提供熱循環測試 (Thermal Cycling) 的允收依據。')
    ]
    for i, (k, v) in enumerate(ipc_mapping):
        ipc_table.rows[i].cells[0].text = k
        ipc_table.rows[i].cells[1].text = v

    # --- 4. 流程管控與準備 (Process Control) ---
    doc.add_heading('4. 流程管控與回溫環境標準', level=1)
    flow = [
        'IQC 進料檢驗：核對保存期限、批次號、粘度檢驗報告 (CoA)。',
        '冷凍存儲：-15°C 至 -40°C。嚴禁儲存在非低溫環境超過 1 小時。',
        '垂直回溫：取出後針頭朝下垂直靜置 1-2 小時。嚴禁用吹風機或烘箱強行回溫。',
        'PCB 預烤：建議 100°C / 1.5 小時以移除水份，防止固化過程中產生大氣泡。',
        '板材預熱：點膠前將 PCB 加熱至 75°C - 90°C，可使膠材流動速度提升 3-5 倍。'
    ]
    for f in flow: doc.add_paragraph(f, style='List Bullet')

    # --- 5. 點膠參數設計與路徑 (Engineering Parameters) ---
    doc.add_heading('5. 點膠參數設計與路徑優化', level=1)
    
    if os.path.exists(IMG['NEEDLE']):
        doc.add_picture(IMG['NEEDLE'], width=Inches(4.5))
        doc.add_paragraph('圖 1: 針頭與元件邊距、點膠高度設定基準 (H=0.3mm, D=0.8mm)', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if os.path.exists(IMG['PATTERNS']):
        doc.add_picture(IMG['PATTERNS'], width=Inches(4.5))
        doc.add_paragraph('圖 2: 推薦 L 型路徑 (引導空氣對角排除，氣鎖率最低)', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

    param_table = doc.add_table(rows=4, cols=2)
    param_table.style = 'Table Grid'
    p_data = [
        ('針頭規格', '23G (Standard) - 25G (Fine Pitch) / 鋼座斜口針'),
        ('點膠壓力', '手動：0.1-0.2 MPa / 自動閥：0.2-0.35 MPa'),
        ('路徑模式', 'L型 (首選)、I型 (適用小型 CSP)、U型 (巨型元件)'),
        ('點完停頓 (Delay)', '建議 150-500ms，確保點膠末端膠滴完全斷開，防止拉絲。')
    ]
    for i, (k, v) in enumerate(p_data):
        param_table.rows[i].cells[0].text = k
        param_table.rows[i].cells[1].text = v

    # --- 6. 應用元件矩陣細則 (Application Matrix) ---
    doc.add_heading('6. 應用元件類別與物理尺寸矩陣對應表', level=1)
    matrix_table = doc.add_table(rows=6, cols=5)
    matrix_table.style = 'Table Grid'
    m_h = ['元件類別', 'Die Size (mm)', 'Ball Pitch (mm)', 'Ball Size (mm)', '工藝應用建議']
    for i, v in enumerate(m_h): matrix_table.rows[0].cells[i].text = v
    m_data = [
        ('Flip Chip', '2x2~15x15', '0.10~0.25', '0.05~0.15', '必須底填 (底部間隙僅 30-80um)'),
        ('Fine Pitch BGA', '5x5~12x12', '0.35~0.50', '0.20~0.30', '高度建議 (防止熱脹冷縮脫焊)'),
        ('WLCSP', '1.5x1.5~8x8', '0.30~0.50', '0.15~0.25', '強制點膠 (防止跌落衝擊)'),
        ('Standard BGA', '15x15~45x45', '0.80~1.27', '0.35~0.60', '建議點膠 (提升大件抗扭力強度)'),
        ('QFN / LGA', '< 10x10', '0.40~0.65', 'N/A (Pad)', '條件適用 (需確認 Standoff 排氣)')
    ]
    for i, r_row in enumerate(m_data):
        for j, val in enumerate(r_row): matrix_table.rows[i+1].cells[j].text = val

    # --- 7. 品質檢驗準則與判定 (Quality Criteria) ---
    doc.add_heading('7. 品質檢驗允收標準 (Acceptance Criteria)', level=1)
    
    # Fillet Section
    doc.add_heading('7.1 Fillet (外溢膠) 目視標準', level=2)
    if os.path.exists(IMG['FILLET']):
        doc.add_picture(IMG['FILLET'], width=Inches(4.2))
        doc.add_paragraph('圖 3: Fillet 高度合格帶 (50% - 75% 元件側厚)', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Voiding Section
    doc.add_heading('7.2 內部空洞 (Voiding) 與分層 (Delamination)', level=2)
    c1, c2 = doc.add_table(rows=1, cols=2).rows[0].cells
    if os.path.exists(IMG['XRAY']):
        p = c1.add_paragraph()
        p.add_run().add_picture(IMG['XRAY'], width=Inches(2.5))
        c1.add_paragraph('圖 4: X-Ray 下的空洞判定 (< 25%)')
    if os.path.exists(IMG['CSAM']):
        p = c2.add_paragraph()
        p.add_run().add_picture(IMG['CSAM'], width=Inches(2.5))
        c2.add_paragraph('圖 5: CSAM 下的分層缺陷比對')

    quality_points = [
        '允收：溢膠寬度 (Fillet Width) 應為元件邊長的 10% 以上，但不超過 1.5mm。',
        '允收：Fillet 圍繞元件周長 100% 完整，無斷裂或嚴重內陷。',
        '拒收：膠材爬升至元件頂部表面 (污染散熱鰭片接觸面)。',
        '拒收：底填膠內部大空洞與相鄰兩錫球同時接觸 (Risk of bridge)。',
        '拒收：元件發生超過 0.1mm 的非預期偏移或抬升。'
    ]
    for q in quality_points: doc.add_paragraph(q, style='List Bullet')

    # --- 8. 異常預防與對策 (Troubleshooting) ---
    doc.add_heading('8. 異常預防與故障排除對策表', level=1)
    ts_table = doc.add_table(rows=6, cols=3)
    ts_table.style = 'Table Grid'
    ts_h = ['異常現象 (Problem)', '潛在根因 (Root Cause)', '工藝改善措施 (Countermeasure)']
    for i, v in enumerate(ts_h): ts_table.rows[0].cells[i].text = v
    ts_data = [
        ('滲透緩慢 / 不全', 'PCB 溫度不足或間隙過窄', '校準基板預熱至 85C；更換更高流速膠材。'),
        ('固化後大量微氣泡', 'PCB 中殘餘水份或回溫不足', '確認 PCB 預烤流程；嚴格控管 2H 垂直回溫。'),
        ('出膠量不穩', '控制器壓力波動或針頭堵塞', '清理計量泵系統；過濾壓縮空氣中的油水。'),
        ('Fillet 表層收縮', '固化溫度上升過快', '調整烘箱 Profile 平緩 Ramping Rate。'),
        ('元件底部空洞過大', '點膠路徑氣鎖 (Air Lock)', '由 U型改為 L型路徑；降低一次點膠量改為二次點膠。')
    ]
    for i, row in enumerate(ts_data):
        for j, val in enumerate(row): ts_table.rows[i+1].cells[j].text = val

    # --- 9. 受控返工程序 (Rework Procedure) ---
    doc.add_heading('9. 受控返工程序 (Rework Procedure)', level=1)
    rework_steps = [
        'A. 預加熱：將組件放入加熱台，預熱全板至 150°C (120 秒)。',
        'B. 拆卸加熱：使用熱風槍或 BGA 返修台，將目標元件加熱至 240°C-250°C。',
        'C. 移除動作：使用真空吸嘴在膠材軟化時取出元件。嚴禁在低溫下施力撬起。',
        'D. 殘膠清理：在 200°C 餘熱下，使用專用聚四氟乙烯刮刀移除殘餘環氧樹脂。',
        'E. 最終清潔：使用異丙基酒精 (IPA) 或原廠推薦環氧樹脂清洗劑進行濕潤清理。'
    ]
    for r in rework_steps: doc.add_paragraph(r, style='List Bullet')

    # Save
    save_path = 'Underfill_Ultimate_Standard_V10_Consolidated.docx'
    doc.save(save_path)
    print(f"Truly Ultimate Document saved to: {os.path.abspath(save_path)}")

if __name__ == '__main__':
    build_true_ultimate_wi(
        r'C:\Users\user\Desktop\Underfill\Underfill_Ultimate_Standard_V10_Consolidated.docx',
        r'C:\Users\user\Desktop\Underfill\extracted_images'
    )
