import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_final_merged_wi():
    doc = Document()
    
    # Global Font Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

    # --- Title Section ---
    title = doc.add_heading('Underfill 全工藝作業標準說明書 (Comprehensive WI)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table_info = doc.add_table(rows=4, cols=4)
    table_info.style = 'Table Grid'
    info = [
        ['文件編號', 'WI-SMT-UF3808-TOTAL', '版本', 'V3.0 (Final)'],
        ['膠材型號', 'Loctite ECCOBOND UF 3808', '應用', 'CSP/BGA/FlipChip'],
        ['制訂部門', 'SMT 工程部 / 品質管理部', '日期', '2026-04-09'],
        ['狀態', '正式發行', '機密性', 'Confidential']
    ]
    for i, row_data in enumerate(info):
        for j, val in enumerate(row_data):
            table_info.rows[i].cells[j].text = val

    doc.add_paragraph('\n')

    # --- 1. 目的與適用範圍 ---
    doc.add_heading('1. 目的與適用範圍', level=1)
    doc.add_paragraph('本文件整合了 Loctite UF 3808 的原廠化學物理特性、精密點膠路徑設計、多維度元件矩陣對應及 IPC 國際品質允收標準。適用於 SMT 產線所有涉及封裝底填之製程管控，旨在提升產品在衝擊、跌落與熱循環下的可靠性。')

    # --- 2. 膠材技術參數 (Henkel Loctite UF 3808) ---
    doc.add_heading('2. 膠材技術參數與原廠規範', level=1)
    table_specs = doc.add_table(rows=11, cols=2)
    table_specs.style = 'Table Grid'
    specs = [
        ('化學基礎 / 固化方式', '單組份環氧樹脂 / 熱固化'),
        ('粘度 (Viscosity)', '348 mPa·s (cP) @ 25°C - 具備高速毛細流動性'),
        ('儲存週期 (Shelf Life)', '12 個月 (@ -40°C 至 -15°C)'),
        ('室溫工作壽命 (Pot Life)', '24 小時 (@ 25°C)'),
        ('玻璃轉化溫度 (Tg)', '113°C (by TMA)'),
        ('CTE (Below Tg / Above Tg)', '55 ppm/°C / 171 ppm/°C'),
        ('彈性模量 (Modulus)', '2.8 GPa (@ 25°C)'),
        ('硬度 (Shore D)', '88'),
        ('標準固化條件', '130°C / 8 mins 或 150°C / 5 mins'),
        ('吸濕率 (Moisture Absorption)', '低於 0.5% (符合高可靠度電子要求)'),
        ('可返修性 (Reworkable)', '是 (加熱至 240°C 以上可移除)')
    ]
    for i, (k, v) in enumerate(specs):
        table_specs.rows[i].cells[0].text = k
        table_specs.rows[i].cells[1].text = v

    # --- 3. 流程與環境管控 ---
    doc.add_heading('3. 流程與環境管控標準', level=1)
    flow_steps = [
        'A. 冷凍存儲：專用冷凍櫃，溫度維持在 -15°C 以下，並落實 FIFO 管理。',
        'B. 解凍流程：斜放或垂直站立於 25°C 環境 60-120 分鐘。嚴禁用強制加熱方式解凍。',
        'C. PCB 預熱：為確保膠材流動性，板材表面溫度建議預熱至 70°C - 90°C。',
        'D. 真空除泡：若有手動裝填需求，須經由離心脫泡機排除剩餘微小氣泡。'
    ]
    for step in flow_steps:
        doc.add_paragraph(step, style='List Bullet')

    # --- 4. 點膠路徑與工藝參數 ---
    doc.add_heading('4. 點膠路徑設計與工藝參數', level=1)
    table_path = doc.add_table(rows=4, cols=2)
    table_path.style = 'Table Grid'
    paths = [
        ('路徑模式', '適用場景說明'),
        ('I型 (Single Edge)', '沿單側長邊點膠。適合 < 10mm 元件。'),
        ('L型 (Adjacent Edges)', '最強推薦模式。膠材由兩側向中心推進，空氣向 135度方向排出，空洞率最低。'),
        ('U型 (Three Edges)', '適合 > 25mm 巨型元件。需注意點跡結合處易產生殘餘空氣。')
    ]
    for i, r in enumerate(paths):
        table_path.rows[i].cells[0].text = r[0]
        table_path.rows[i].cells[1].text = r[1]
    
    p = doc.add_paragraph('\n建議參數設定：')
    p.add_run('\n- 針頭規格: 23G 至 25G (依間隙 Stand-off 調整)')
    p.add_run('\n- 點膠壓力: 0.1Mpa - 0.35Mpa')
    p.add_run('\n- 點膠高度: 0.2mm - 0.5mm (視元件厚度)')

    # --- 5. 元件類別尺寸矩陣表 ---
    doc.add_heading('5. 元件類別尺寸矩陣對應表', level=1)
    table_matrix = doc.add_table(rows=6, cols=5)
    table_matrix.style = 'Table Grid'
    h_matrix = ['元件類型', 'Die Size (mm)', 'Ball Pitch (mm)', 'Ball Size (mm)', '應用判定']
    for i, v in enumerate(h_matrix):
        table_matrix.rows[0].cells[i].text = v
    m_data = [
        ('BGA (Big)', '25~45', '0.8~1.27', '0.45~0.60', '推薦 (機械增強)'),
        ('BGA (Small)', '10~25', '0.5~0.8', '0.30~0.45', '推薦 (可靠度)'),
        ('WLCSP/CSP', '1.5~10', '0.3~0.5', '0.15~0.25', '強制 (跌落測試)'),
        ('Flip Chip', '2~15', '0.15~0.3', '0.05~0.15', '核心應用 (底部封裝)'),
        ('QFN/LGA', '< 12', '0.4~0.65', 'N/A', '條件適用 (評估排氣)')
    ]
    for i, row in enumerate(m_data):
        for j, val in enumerate(row):
            table_matrix.rows[i+1].cells[j].text = val

    # --- 6. 品質檢驗標準 (依據 IPC 規範) ---
    doc.add_heading('6. 品質檢驗標準與 IPC 規範', level=1)
    doc.add_heading('6.1 IPC 專項細則', level=2)
    table_ipc = doc.add_table(rows=4, cols=2)
    table_ipc.style = 'Table Grid'
    ipc_info = [
        ('IPC-J-STD-030', '底填膠填充定義：必須填滿元件下方所有空間，不得有貫穿性裂紋。'),
        ('IPC-A-610 Section 8.3', '外觀允收：包封 (Fillet) 寬度須超過元件邊緣，高度覆蓋率 50%~100%。'),
        ('IPC-7095 Section 7.5', '空洞判定：單一焊點空洞面積 < 25%，總空洞區域 < 總面積之 25%。')
    ]
    for i, r in enumerate(ipc_info):
        table_ipc.rows[i].cells[0].text = r[0]
        table_ipc.rows[i].cells[1].text = r[1]

    doc.add_heading('6.2 判定準則 (Accept / Reject)', level=2)
    criteria = [
        '允收：Fillet 完整包裹四周元件腳位，高度達元件側壁 2/3。',
        '允收：表面平整、無肉眼可見的大型空洞或斷膠。',
        '拒收：膠材爬升至元件頂部表面 (Top Surface Contamination)。',
        '拒收：膠材溢出範位超過周邊 1.5mm，或污染鄰近測試點。',
        '拒收：元件發生超過 0.1mm 的明顯水平位移或單邊抬升。'
    ]
    for c in criteria:
        doc.add_paragraph(c, style='List Bullet')

    # --- 7. 故障排除 (Troubleshooting) ---
    doc.add_heading('7. 常見異常原因與對策', level=1)
    table_ts = doc.add_table(rows=5, cols=3)
    table_ts.style = 'Table Grid'
    ts_h = ['異常現象', '可能原因', '排除措施']
    for i, v in enumerate(ts_h):
        table_ts.rows[0].cells[i].text = v
    ts_data = [
        ('點膠不均', '計量泵壓力不穩或針管氣泡', '重新脫泡並校準泵流量參數。'),
        ('滲透緩慢', '基板預熱不足或間隙過小', '調高預熱溫度至 85°C，縮短點膠間距。'),
        ('固化後分層', 'PCB 表面有殘餘助焊劑', '加強回流焊後的清洗流程或檢查助焊劑兼容性。'),
        ('Fillet 氣泡', '膠材回溫不足或 PCB 受潮', '落實回溫時間管理，對 PCB 進行預熱烘烤 (100°C/1h)。')
    ]
    for i, row in enumerate(ts_data):
        for j, val in enumerate(row):
            table_ts.rows[i+1].cells[j].text = val

    # --- 8. 返工程序 (Rework) ---
    doc.add_heading('8. 返工程序 (Rework Procedure)', level=1)
    rework = [
        '1. 預加熱：將組件放置於 BGA 返修台，均勻預熱底部至 150°C。',
        '2. 局部加熱：上方噴嘴加熱至 240°C-250°C (膠材與焊點同時弱化)。',
        '3. 移除：使用真空吸嘴或鑷子移除元件。',
        '4. 清理：在高溫下使用專用刮刀刮除殘餘環氧樹脂，並配合溶劑清理焊盤。'
    ]
    for r in rework:
        doc.add_paragraph(r, style='List Bullet')

    save_path = 'Underfill_WI_Full_Integrated.docx'
    doc.save(save_path)
    print(f'Full Integrated Document saved to: {os.path.abspath(save_path)}')

if __name__ == '__main__':
    create_final_merged_wi()
