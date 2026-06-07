from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def build_ultimate_complete_wi(source_docx, media_dir, out_path):
    # We will build it from scratch with the intended sequence to ensure sections are filled
    doc = Document()
    
    # Global Font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(10)

    # Image Paths
    img_map = {
        'patterns': os.path.join(media_dir, 'word/media/image1.jpeg'),
        'fillet': os.path.join(media_dir, 'word/media/image2.jpeg'),
        'voiding': os.path.join(media_dir, 'word/media/image3.jpeg'),
        'delamination': os.path.join(media_dir, 'word/media/image4.jpeg'),
        'needle': os.path.join(media_dir, 'word/media/image5.jpeg')
    }

    # --- Title ---
    doc.add_heading('Underfill Operator Process Guide - Ultimate Edition', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 1 & 2 (Brief Intro) ---
    doc.add_heading('1. Purpose and Scope', level=1)
    doc.add_paragraph('Standardize the operation of Loctite UF 3808 to ensure board-level reliability.')

    doc.add_heading('2. Source Documents Consolidated', level=1)
    doc.add_paragraph('Consolidated Henkel TDS/SDS and IPC International standards.')

    # --- 3. International Guideline and Specification Mapping (COMPLETED) ---
    doc.add_heading('3. International Guideline and Specification Mapping', level=1)
    ipc_table = doc.add_table(rows=5, cols=2)
    ipc_table.style = 'Table Grid'
    ipc_data = [
        ('IPC-J-STD-030', 'Standard for Selection and Application of Board Level Underfill Materials. Defines flow rates and curing profiles.'),
        ('IPC-A-610', 'Acceptability of Electronic Assemblies. Section 8 defines visual fillet requirements and surface contamination.'),
        ('IPC-7095', 'Design and Assembly Process for BGAs. Specifies that total void area should be < 25% of the solder joint area.'),
        ('IPC-7711/7721', 'Rework and Repair. Defines the thermal profile required (typically > 240°C) for reworkable epoxy removal.'),
        ('IPC-SM-785', 'Guidelines for Accelerated Reliability Testing. Used for defining thermal cycle parameters.')
    ]
    for i, (k, v) in enumerate(ipc_data):
        ipc_table.rows[i].cells[0].text = k
        ipc_table.rows[i].cells[1].text = v

    # --- 4. Material and Storage Specification (UF 3808) (COMPLETED) ---
    doc.add_heading('4. Material and Storage Specification (UF 3808)', level=1)
    doc.add_paragraph('Technical details for Henkel Loctite ECCOBOND UF 3808:')
    spec_table = doc.add_table(rows=8, cols=2)
    spec_table.style = 'Table Grid'
    specs = [
        ('Storage Temperature', '-40°C to -15°C (Shelf Life: 12 Months)'),
        ('Thawing Time', '60 - 120 Minutes (Vertical orientation, needle down)'),
        ('Work Life (Pot Life)', '24 Hours @ 25°C room temperature'),
        ('Viscosity', '348 mPa·s (cP) @ 25°C, 1,000 s⁻¹'),
        ('Curing Profile', '130°C for 8 minutes OR 150°C for 5 minutes'),
        ('Tg (by TMA)', '113°C'),
        ('CTE (Below / Above Tg)', '55 ppm/°C / 171 ppm/°C'),
        ('Storage Modulus', '2.8 GPa @ 25°C')
    ]
    for i, (k, v) in enumerate(specs):
        spec_table.rows[i].cells[0].text = k
        spec_table.rows[i].cells[1].text = v

    # --- 5. Standard Operator Process Flow ---
    doc.add_heading('5. Standard Operator Process Flow', level=1)
    flow = ["Storage (-20C) -> Thaw (RT, 2h) -> Pre-heat PCB (85C) -> Dispense -> Cure (Oven) -> Inspect"]
    for f in flow: doc.add_paragraph(f, style='List Bullet')

    # --- 6. Dispensing Parameter Window ---
    doc.add_heading('6. Dispensing Parameter Window', level=1)
    doc.add_paragraph('Standard Settings for Needle and Path:')
    if os.path.exists(img_map['needle']):
        doc.add_picture(img_map['needle'], width=Inches(4.5))
    if os.path.exists(img_map['patterns']):
        doc.add_picture(img_map['patterns'], width=Inches(4.5))

    # --- 7. Component-Specific Applicability Matrix (COMPLETED) ---
    doc.add_heading('7. Component-Specific Applicability Matrix', level=1)
    matrix_table = doc.add_table(rows=6, cols=5)
    matrix_table.style = 'Table Grid'
    headers = ['Category', 'Die Size (mm)', 'Ball Pitch (mm)', 'Ball Size (mm)', 'Suitability']
    for i, h in enumerate(headers): matrix_table.rows[0].cells[i].text = h
    m_data = [
        ('BGA', '15~45', '0.65~1.27', '0.35~0.60', 'Recommended (Structural)'),
        ('WLCSP', '1.5~10', '0.30~0.50', '0.15~0.25', 'Mandatory (Drop Test)'),
        ('Flip Chip', '2~12', '0.15~0.30', '0.05~0.15', 'Critical (Gap ≤ 80um)'),
        ('QFN', '< 12', '0.40~0.65', 'N/A', 'Evaluated (Airflow risk)'),
        ('0201 Passives', '< 0.6', 'N/A', 'N/A', 'NOT Recommended')
    ]
    for i, row in enumerate(m_data):
        for j, val in enumerate(row): matrix_table.rows[i+1].cells[j].text = val

    # --- 8. Quality Acceptance Criteria ---
    doc.add_heading('8. Quality Acceptance Criteria (Operator + QA)', level=1)
    if os.path.exists(img_map['fillet']): doc.add_picture(img_map['fillet'], width=Inches(4.0))
    if os.path.exists(img_map['voiding']): doc.add_picture(img_map['voiding'], width=Inches(4.0))
    if os.path.exists(img_map['delamination']): doc.add_picture(img_map['delamination'], width=Inches(4.0))

    # --- 9. Abnormal Handling and Countermeasures (COMPLETED) ---
    doc.add_heading('9. Abnormal Handling and Countermeasures', level=1)
    ts_table = doc.add_table(rows=5, cols=3)
    ts_table.style = 'Table Grid'
    ts_h = ['Defect', 'Possible Root Cause', 'Countermeasure']
    for i, v in enumerate(ts_h): ts_table.rows[0].cells[i].text = v
    ts_data = [
        ('Excessive Voids', 'Entrapped air or PCB moisture', 'Pre-bake PCB at 100°C for 1.5h; Check L-pattern.'),
        ('Bleed-out', 'Over-dispensing or low standoff', 'Reduce pressure or increase robot speed; check gap.'),
        ('Delamination', 'Surface contamination (flux residue)', 'Enhanced cleaning after reflow; verify compatibility.'),
        ('Incomplete Flow', 'Board temperature too low', 'Recalibrate heater plate to ensure 85°C surface temp.')
    ]
    for i, row in enumerate(ts_data):
        for j, val in enumerate(row): ts_table.rows[i+1].cells[j].text = val

    # --- 10 & 11 (Remaining) ---
    doc.add_heading('10. Rework Guideline (Controlled)', level=1)
    doc.add_paragraph('Heat locally to 240°C. Remove component using vacuum. Clean pads while hot.')

    doc.add_heading('11. Operator Checklist (Shift Use)', level=1)
    doc.add_paragraph('[ ] Expiry Check  [ ] Thawing Time (2h)  [ ] Board Temp (85C)')

    doc.save(out_path)
    print(f"Ultimate Completed Document saved to: {out_path}")

if __name__ == "__main__":
    build_ultimate_complete_wi(
        r'C:\Users\user\Desktop\Underfill\Underfill_Operator_Process_Guide_International_Summary.docx',
        r'C:\Users\user\Desktop\Underfill\extracted_images',
        r'C:\Users\user\Desktop\Underfill\Underfill_Final_Official_Complete.docx'
    )
