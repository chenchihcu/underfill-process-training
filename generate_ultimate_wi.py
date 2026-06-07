import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_ultimate_wi():
    doc = Document()
    
    # Image Paths (Absolute)
    IPATHS = {
        'PATTERNS': r'C:\Users\user\.gemini\antigravity\brain\db344bb8-a2bb-486c-81a4-f44245a98563\underfill_dispensing_patterns_1775741805036.png',
        'FILLET': r'C:\Users\user\.gemini\antigravity\brain\db344bb8-a2bb-486c-81a4-f44245a98563\underfill_fillet_height_1775741826978.png',
        'XRAY': r'C:\Users\user\.gemini\antigravity\brain\db344bb8-a2bb-486c-81a4-f44245a98563\underfill_voiding_xray_1775741846746.png',
        'CSAM': r'C:\Users\user\.gemini\antigravity\brain\db344bb8-a2bb-486c-81a4-f44245a98563\csam_underfill_inspection_1775742090770.png',
        'NEEDLE': r'C:\Users\user\.gemini\antigravity\brain\db344bb8-a2bb-486c-81a4-f44245a98563\underfill_needle_setup_1775742113973.png'
    }

    # Font Configuration
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

    # --- Header Table ---
    title = doc.add_heading('Underfill (底填膠) 作業指導與品質標準書', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    header_table = doc.add_table(rows=5, cols=4)
    header_table.style = 'Table Grid'
    hdr = [
        ['文件名稱', 'Underfill Ultimate Operating Standard', '文件編號', 'WI-QA-UF-2026-V5'],
        ['膠材型號', 'Loctite ECCOBOND UF 3808', '製訂日期', '2026-04-09'],
        ['Ball Pitch', '0.1mm - 1.27mm', 'Die Size', '1.5mm - 45.0mm'],
        ['IPC 級別', 'Class 2 / Class 3', '安全等級', 'Confidential'],
        ['狀態', '工藝優化與圖文整合版', '頁數', '1 / 1']
    ]
    for i, r_data in enumerate(hdr):
        for j, val in enumerate(r_data): header_table.rows[i].cells[j].text = val

    # --- 1. 概述與材料規格 (Material Specs) ---
    doc.add_heading('1. 材料規格與特性分析 (Loctite UF 3808)', level=1)
    p = doc.add_paragraph('Loctite UF 3808 是業界標竿的可返修型底填膠，具備極佳的熱循環穩定性。')
    
    spec_table = doc.add_table(rows=10, cols=2)
    spec_table.style = 'Table Grid'
    specs = [
        ('粘度 (Viscosity)', '348 mPa·s - 超高流速，適合細間距 FlipChip'),
        ('固化推薦 (Cure)', '130°C / 8 mins 或 150°C / 5 mins'),
        ('產地與壽命', 'Henkel / 12 Months (@ -40°C)'),
        ('Tg (玻璃轉化點)', '113°C - 確保高溫工作環境下的機械模量'),
        ('CTE 1 / CTE 2', '55 ppm/°C / 171 ppm/°C'),
        ('工作時間 (Work Life)', '24 小時 (@ 25°C - 嚴禁超時使用)'),
        ('彈性模量 (Modulus)', '2.8 GPa - 提供極佳的跌落衝擊墊片效果'),
        ('硬度 (Shore D)', '88'),
        ('吸濕率', '< 0.5% (符合 IPC 高可靠度產品)'),
        ('化學抵抗', '通過各種清洗劑與助焊劑兼容性測試')
    ]
    for i, (k, v) in enumerate(specs):
        spec_table.rows[i].cells[0].text = k
        spec_table.rows[i].cells[1].text = v

    # --- 2. 設備與工藝準備 (Process Preparation) ---
    doc.add_heading('2. 設備設定與點膠參數', level=1)
    doc.add_paragraph('精密點膠需要嚴格控制針頭位置與工作壓力。下圖為標準針頭設定基準。')
    
    if os.path.exists(IPATHS['NEEDLE']):
        doc.add_picture(IPATHS['NEEDLE'], width=Inches(4.5))
        doc.add_paragraph('圖 1: 點膠針頭與元件邊距、高度設定基準', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    param_list = [
        '針頭規格：建議 23G (普通 BGA) 或 25G-27G (Fine Pitch)。',
        '點膠壓力：手動泵 0.1MPa / 自動閥 0.25MPa。',
        '基板預熱：70°C - 90°C (必須！確保膠材能均勻滲透中心區)。',
        '針頭高度：元件頂部下 0.2mm - 0.5mm 處。'
    ]
    for item in param_list: doc.add_paragraph(item, style='List Bullet')

    # --- 3. 點膠路徑與流動模式 (Patterns) ---
    doc.add_heading('3. 點膠路徑設計與流動引導', level=1)
    doc.add_paragraph('依據元件尺寸選擇路徑，核心目標是「空氣完全排出」。')
    
    if os.path.exists(IPATHS['PATTERNS']):
        doc.add_picture(IPATHS['PATTERNS'], width=Inches(4.5))
        doc.add_paragraph('圖 2: L型推薦點膠路徑示意圖 (兩側推進，對角排氣)', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('3.1 元件尺寸對應矩陣表', level=2)
    matrix_table = doc.add_table(rows=6, cols=5)
    matrix_table.style = 'Table Grid'
    m_h = ['類別', 'Die Size (mm)', 'Pitch (mm)', 'Ball Size', '判定說明']
    for i, v in enumerate(m_h): matrix_table.rows[0].cells[i].text = v
    m_data = [
        ('BGA', '15~45', '0.65~1.27', '0.35~0.60', '推薦 (提升重摔強度)'),
        ('WLCSP', '1.5~8', '0.30~0.50', '0.15~0.25', '強制 (防止焊點裂紋)'),
        ('FlipChip', '2~12', '0.15~0.30', '0.05~0.15', '核心 (填補 50um 間隙)'),
        ('QFN', '< 12', '0.40~0.65', 'N/A', '條件適用 (評估排氣)'),
        ('Passive', '< 0402', 'N/A', 'N/A', '不建議使用 (易致偏位)')
    ]
    for i, r_data in enumerate(m_data):
        for j, val in enumerate(r_data): matrix_table.rows[i+1].cells[j].text = val

    # --- 4. 品質判定標準 (Quality & IPC) ---
    doc.add_heading('4. 品質允收標準與檢驗規範 (IPC Standards)', level=1)
    
    doc.add_heading('4.1 Fillet (膠側) 視覺判定', level=2)
    doc.add_paragraph('側邊包封 (Fillet) 應連續且飽滿，高度依據 IPC-A-610 規範定義。')
    
    if os.path.exists(IPATHS['FILLET']):
        doc.add_picture(IPATHS['FILLET'], width=Inches(4.2))
        doc.add_paragraph('圖 3: Fillet 高度合格標準 (50% - 75% Component Height)', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('4.2 內部空洞 (Voiding) - X-Ray 判定', level=2)
    if os.path.exists(IPATHS['XRAY']):
        doc.add_picture(IPATHS['XRAY'], width=Inches(4.0))
        doc.add_paragraph('圖 4: X-Ray 掃描空洞判定 (總空洞面積 < 25%)', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('4.3 分層檢測 (Delamination) - CSAM 檢驗', level=2)
    doc.add_paragraph('對於高可靠度產品 (Class 3)，須進行超音波掃描以偵測界面分層。')
    
    if os.path.exists(IPATHS['CSAM']):
        doc.add_picture(IPATHS['CSAM'], width=Inches(4.2))
        doc.add_paragraph('圖 5: CSAM 檢驗合格對比 (對比分層缺陷)', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 5. 異常對策與返工 (Troubleshooting & Rework) ---
    doc.add_heading('5. 異常分析與返修規範', level=1)
    
    ts_table = doc.add_table(rows=5, cols=3)
    ts_table.style = 'Table Grid'
    ts_data = [
        ('異常描述', '潛在根因 (Root Cause)', '對策 (Countermeasure)'),
        ('Fillet 不連續', '膠材回溫不足或雜質阻塞', '嚴嚴落實 2H 回溫；更換過濾針頭。'),
        ('內部大氣泡', 'PCB 未預烤 (Moisture)', '100°C / 1.5H 預烤排除水氣。'),
        ('組件偏位', '點膠壓力過大或間距太近', '調降點膠初壓力；增加點膠避讓間距。'),
        ('滲透不全', '基板預熱偏差', '校準加熱台溫控，確保 PCB 達 80°C。')
    ]
    for i, r_data in enumerate(ts_data):
        for j, val in enumerate(r_data): ts_table.rows[i].cells[j].text = val

    save_path = 'Underfill_Ultimate_Standard_V3.docx'
    doc.save(save_path)
    print(f'Ultimate Document saved to: {os.path.abspath(save_path)}')

if __name__ == '__main__':
    create_ultimate_wi()
