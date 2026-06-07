from docx import Document
from docx.shared import Inches
import os

def build_final_relocated(source_docx, media_dir, out_path):
    doc = Document(source_docx)
    new_doc = Document()
    
    # Map extracted images to topics (heuristic based on order they appeared in appendix)
    # image1-5 likely correspond to the 5 visuals generated earlier in that order
    img_map = {
        'patterns': os.path.join(media_dir, 'word/media/image1.jpeg'),
        'fillet': os.path.join(media_dir, 'word/media/image2.jpeg'),
        'voiding': os.path.join(media_dir, 'word/media/image3.jpeg'),
        'delamination': os.path.join(media_dir, 'word/media/image4.jpeg'),
        'needle': os.path.join(media_dir, 'word/media/image5.jpeg')
    }

    # Find headers
    headers = {}
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if p.style.name.startswith('Heading 1'):
            if "5. Standard Operator Process Flow" in text: headers['flow'] = i
            if "6. Dispensing Parameter Window" in text: headers['params'] = i
            if "8. Quality Acceptance Criteria" in text: headers['quality'] = i
            if "10. Rework Guideline" in text: headers['rework'] = i
            if "12. Figure Appendix" in text: headers['appendix'] = i

    for i, p in enumerate(doc.paragraphs):
        if i >= headers.get('appendix', 999):
            break
            
        new_p = new_doc.add_paragraph(p.text, style=p.style)
        
        # Insert images after specific headers
        if "6. Dispensing Parameter Window" in p.text:
            if os.path.exists(img_map['patterns']):
                new_doc.add_paragraph("視覺圖解：點膠路徑 (L型路徑推薦)").bold = True
                new_doc.add_picture(img_map['patterns'], width=Inches(4.5))
            if os.path.exists(img_map['needle']):
                new_doc.add_paragraph("視覺圖解：針頭與針筒設定基準").bold = True
                new_doc.add_picture(img_map['needle'], width=Inches(4.5))
                
        if "8. Quality Acceptance Criteria" in p.text:
            if os.path.exists(img_map['fillet']):
                new_doc.add_paragraph("品質標準：Fillet 高度要求 (50-75%)").bold = True
                new_doc.add_picture(img_map['fillet'], width=Inches(4.0))
            if os.path.exists(img_map['voiding']):
                new_doc.add_paragraph("品質標準：內部空洞 (X-Ray 判定)").bold = True
                new_doc.add_picture(img_map['voiding'], width=Inches(4.0))
            if os.path.exists(img_map['delamination']):
                new_doc.add_paragraph("品質標準：分層與剝離 (CSAM 檢驗)").bold = True
                new_doc.add_picture(img_map['delamination'], width=Inches(4.0))

    new_doc.save(out_path)
    print(f"Relocated photos document saved to: {out_path}")

if __name__ == "__main__":
    build_final_relocated(
        r'C:\Users\user\Desktop\Underfill\Underfill_Operator_Process_Guide_International_Summary.docx',
        r'C:\Users\user\Desktop\Underfill\extracted_images',
        r'C:\Users\user\Desktop\Underfill\Underfill_Operator_Process_Guide_Final_Relocated.docx'
    )

