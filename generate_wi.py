import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_wi():
    doc = Document()
    
    # Set default font to Microsoft JhengHei if possible, or a standard gothic/sans
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(11)
    # Important for Chinese characters in some docx readers
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

    # --- Title ---
    title = doc.add_heading('Underfill 作業指導書 (Standard Operating Instruction)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Document Info ---
    table_info = doc.add_table(rows=3, cols=4)
    table_info.style = 'Table Grid'
    cells = table_info.rows[0].cells
    cells[0].text = '文件編號'
    cells[1].text = 'WI-SMT-UF001'
    cells[2].text = '文件版本'
    cells[3].text = 'V1.0'
    
    cells = table_info.rows[1].cells
    cells[0].text = '膠材型號'
    cells[1].text = 'Loctite UF 3808'
    cells[2].text = '發行日期'
    cells[3].text = '2026-04-09'
    
    cells = table_info.rows[2].cells
    cells[0].text = '制訂人員'
    cells[1].text = 'Engineering Dept.'
    cells[2].text = '頁數'
    cells[3].text = '1 / 1'

    doc.add_paragraph('\n')

    # --- Section 1: 目的 ---
    doc.add_heading('1. 目的', level=1)
    doc.add_paragraph('規範 Loctite ECCOBOND UF 3808 底填膠的標準作業流程，確保 BGA/CSP 元件在組裝後的機械強度與熱可靠度，防止跌落、衝撃或熱循環導致的焊點裂紋。')

    # --- Section 2: 膠材規格與特性 ---
    doc.add_heading('2. 膠材規格與特性 (Loctite UF 3808)', level=1)
    p = doc.add_paragraph()
    p.add_run('UF 3808 是一款專為移動通訊設計的高流速、可返修快乾型封裝底填膠。').bold = True
    
    table_spec = doc.add_table(rows=6, cols=2)
    table_spec.style = 'Table Grid'
    spec_data = [
        ('項目', '技術指標 / 規格說明'),
        ('化學類型', '單組份環氧樹脂 (One-component Epoxy)'),
        ('粘度 (Viscosity)', '348 cP (@ 25°C, 1,000 s⁻¹) - 極快流速'),
        ('固化條件', '130°C / 8 mins 或 150°C / 5 mins'),
        ('玻璃轉化溫度 (Tg)', '113°C (TMA)'),
        ('熱膨脹係數 (CTE)', '55 ppm/°C (<Tg) / 171 ppm/°C (>Tg)')
    ]
    for i, (k, v) in enumerate(spec_data):
        table_spec.rows[i].cells[0].text = k
        table_spec.rows[i].cells[1].text = v

    # --- Section 3: 流程控管 (Process Flow) ---
    doc.add_heading('3. 標準流程控管', level=1)
    flow = [
        '1. 冷凍存儲: -15°C 至 -40°C，嚴格控管先進先出 (FIFO)。',
        '2. 常溫回溫: 針頭朝下垂直放置 1-2 小時，嚴禁用烘箱加熱。',
        '3. PCB 預烤: 100°C / 2 hrs (選配)，移除板材水份防止固化空洞。',
        '4. 針管預熱: 選配，若流速慢可設定加熱套至 30-40°C。',
        '5. 精準點膠: 依標準路徑進行，監控膠材重量與流出寬度。',
        '6. 烘烤固化: 輸送帶烘箱，實測板面溫度須達 130-150°C 區間。',
        '7. 品質檢驗: 100% 目視檢查及抽樣 X-Ray 檢測。'
    ]
    for item in flow:
        doc.add_paragraph(item, style='List Bullet')

    # --- Section 4: 點膠工藝參數與路徑 ---
    doc.add_heading('4. 點膠工藝參數與路徑', level=1)
    doc.add_paragraph('良好的底填效果取決於「空氣排出」的路徑設計。')
    
    # Dispensing Patterns Table
    table_pattern = doc.add_table(rows=4, cols=2)
    table_pattern.style = 'Table Grid'
    pattern_data = [
        ('路徑名稱', '說明'),
        ('I型 (Single Side)', '僅在長邊單側點膠。適用於 10x10mm 以下的小型元件。'),
        ('L型 (Two Sides)', '最推薦路徑。膠材從兩側推進，空氣由對角處排出。'),
        ('U型 (Three Sides)', '適用於 25x25mm 以上大型元件，但需控制點膠延遲，避免包圍空洞。')
    ]
    for i, (n, d) in enumerate(pattern_data):
        table_pattern.rows[i].cells[0].text = n
        table_pattern.rows[i].cells[1].text = d

    # Parameters
    p = doc.add_paragraph('\n建議參數設置：')
    p.add_run('\n- 針頭規格: 23G - 25G (依元件間隙調整)')
    p.add_run('\n- 點膠壓力: 0.1 - 0.3 MPa')
    p.add_run('\n- 基板預熱: 70 - 90°C (提升浸潤速度)')
    p.add_run('\n- 針頭距離: 元件邊緣 0.5 - 0.8 mm')

    # --- Section 5: 應用元件矩陣 (Matrix) ---
    doc.add_heading('5. 應用元件類別矩陣表', level=1)
    table_matrix = doc.add_table(rows=5, cols=3)
    table_matrix.style = 'Table Grid'
    matrix_headers = ['元件類別', '適用性', '技術備註']
    for i, h in enumerate(matrix_headers):
        table_matrix.rows[0].cells[i].text = h
    
    matrix_data = [
        ('BGA / CSP / Pop', '推薦 (O)', '標準流程，確保結構強化。'),
        ('WLCSP / Flip Chip', '強制 (O)', '必須使用，以抵消熱應力失配。'),
        ('QFN / MLF', '條件適用 (△)', '需 Standoff > 0.1mm，避免膠材被氣體推開。'),
        ('LGA / Shielding', '評估 (X)', '間隙太低或範圍太大，容易產生大量空洞。')
    ]
    for i, row in enumerate(matrix_data):
        for j, val in enumerate(row):
            table_matrix.rows[i+1].cells[j].text = val

    # --- Section 6: 品質檢驗標準與 IPC 規範 ---
    doc.add_heading('6. 品質檢驗標準 (Quality Standards)', level=1)
    doc.add_paragraph('依據 IPC 國際通用規範進行判定：')
    
    ipc_list = [
        'IPC-J-STD-030: 核心標準，詳述 Underfill 點膠、流動與固化的工藝規範。',
        'IPC-A-610: 針對 SMT 元件的外觀接受性標準（斷裂、溢流、污染）。',
        'IPC-7095: BGA 設計與組裝，規定內部空洞總面積需 < 25%。',
        'IPC-7711/7721: 規範底填膠元件的返修流程（加熱至 240°C+ 進行移除）。'
    ]
    for ipc in ipc_list:
        doc.add_paragraph(ipc, style='List Bullet')

    # Acceptance Criteria
    p = doc.add_paragraph('允收標準：', style='Normal')
    p.add_run('\n1. 溢膠包封 (Fillet) 高度應介於元件厚度的 50%-75% 之間。')
    p.add_run('\n2. 元件頂部表面不可有任何膠材污染。')
    p.add_run('\n3. X-Ray 抽測單顆錫球下方的空洞不可超過該錫球投影面積的 25%。')

    # --- Section 7: 異常對策 (Troubleshooting) ---
    doc.add_heading('7. 常見異常與對策', level=1)
    trouble = doc.add_table(rows=4, cols=3)
    trouble.style = 'Table Grid'
    trouble.rows[0].cells[0].text = '異常現象'
    trouble.rows[0].cells[1].text = '可能原因'
    trouble.rows[0].cells[2].text = '改善措施'
    
    t_data = [
        ('大量空洞 (Voids)', '板材潮濕或路徑氣鎖', 'PCB 增加預烤；改用 L 型點膠。'),
        ('滲透不全', '基板溫度太低或 Standoff 太低', '增加基板預熱至 85°C；檢查元件規格。'),
        ('溢流污染 (Bleed-out)', '點膠量過多或壓力不穩', '降低出膠量；清理針頭污漬。')
    ]
    for i, row in enumerate(t_data):
        for j, val in enumerate(row):
            trouble.rows[i+1].cells[j].text = val

    save_path = 'Underfill_WI_UF3808_Enhanced.docx'
    doc.save(save_path)
    print(f'Document saved to: {os.path.abspath(save_path)}')

if __name__ == '__main__':
    create_wi()
