from docx import Document
from docx.shared import Inches
import os

def rearrange_images(file_path, output_path):
    doc = Document(file_path)
    
    # Identify the section headers
    headers = {}
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if p.style.name.startswith('Heading 1'):
            if "5. Standard Operator Process Flow" in text: headers['flow'] = i
            if "6. Dispensing Parameter Window" in text: headers['params'] = i
            if "8. Quality Acceptance Criteria" in text: headers['quality'] = i
            if "10. Rework Guideline" in text: headers['rework'] = i
            if "12. Figure Appendix" in text: headers['appendix'] = i

    # Images I generated previously (absolute paths)
    IMAGES = {
        'PATTERNS': r'C:\Users\user\.gemini\antigravity\brain\db344bb8-a2bb-486c-81a4-f44245a98563\underfill_dispensing_patterns_1775741805036.png',
        'FILLET': r'C:\Users\user\.gemini\antigravity\brain\db344bb8-a2bb-486c-81a4-f44245a98563\underfill_fillet_height_1775741826978.png',
        'XRAY': r'C:\Users\user\.gemini\antigravity\brain\db344bb8-a2bb-486c-81a4-f44245a98563\underfill_voiding_xray_1775741846746.png',
        'CSAM': r'C:\Users\user\.gemini\antigravity\brain\db344bb8-a2bb-486c-81a4-f44245a98563\csam_underfill_inspection_1775742090770.png',
        'NEEDLE': r'C:\Users\user\.gemini\antigravity\brain\db344bb8-a2bb-486c-81a4-f44245a98563\underfill_needle_setup_1775742113973.png'
    }

    # Create a new document to reconstruct
    new_doc = Document()
    
    # We will iterate through paragraphs of the original and insert images at key points
    for i, p in enumerate(doc.paragraphs):
        # Skip the appendix section entirely in the new doc
        if i >= headers.get('appendix', 999):
            break
            
        # Copy paragraph
        new_p = new_doc.add_paragraph(p.text, style=p.style)
        
        # Insert images after specific headers
        if "6. Dispensing Parameter Window" in p.text:
            if os.path.exists(IMAGES['PATTERNS']):
                new_doc.add_paragraph("Figure: Dispensing Path (L-Pattern Recommended)").bold = True
                new_doc.add_picture(IMAGES['PATTERNS'], width=Inches(4.5))
            if os.path.exists(IMAGES['NEEDLE']):
                new_doc.add_paragraph("Figure: Needle and Syringe Setup").bold = True
                new_doc.add_picture(IMAGES['NEEDLE'], width=Inches(4.5))
                
        if "8. Quality Acceptance Criteria" in p.text:
            if os.path.exists(IMAGES['FILLET']):
                new_doc.add_paragraph("Figure: Fillet Height Criteria (50-75%)").bold = True
                new_doc.add_picture(IMAGES['FILLET'], width=Inches(4.0))
            if os.path.exists(IMAGES['XRAY']):
                new_doc.add_paragraph("Figure: Underfill Voiding (X-Ray Check)").bold = True
                new_doc.add_picture(IMAGES['XRAY'], width=Inches(4.0))
            if os.path.exists(IMAGES['CSAM']):
                new_doc.add_paragraph("Figure: Delamination Check (CSAM)").bold = True
                new_doc.add_picture(IMAGES['CSAM'], width=Inches(4.0))

    new_doc.save(output_path)
    print(f"Relocated images document saved to: {output_path}")

if __name__ == "__main__":
    rearrange_images(
        r'C:\Users\user\Desktop\Underfill\Underfill_Operator_Process_Guide_International_Summary.docx',
        r'C:\Users\user\Desktop\Underfill\Underfill_Operator_Process_Guide_Integrated_Visuals.docx'
    )
