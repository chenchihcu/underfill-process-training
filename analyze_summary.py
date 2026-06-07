from docx import Document
import os

def analyze_doc(file_path):
    doc = Document(file_path)
    print(f"Analyzing: {file_path}")
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    
    # Check for images in paragraphs or document body
    # python-docx doesn't easily show images in-line in a simple list without iterating xml or shapes
    # But we can check for inline shapes
    inline_shapes = doc.inline_shapes
    print(f"Total Inline Shapes (Images): {len(inline_shapes)}")
    
    # List headers
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading') or para.text.isupper():
            if len(para.text.strip()) > 3:
                print(f"P[{i}] Style: {para.style.name} | Text: {para.text[:50]}...")

if __name__ == "__main__":
    analyze_doc(r'C:\Users\user\Desktop\Underfill\Underfill_Operator_Process_Guide_International_Summary.docx')
